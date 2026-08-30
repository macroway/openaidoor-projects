"""第六章：PDF、DOCX、XLSX 的解析、清洗与切分轻量示例。"""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

import pdfplumber
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


SAMPLE_FILES = {
    "pdf": "恒温设备维护手册.pdf",
    "docx": "客户服务处理制度.docx",
    "xlsx": "产品故障案例.xlsx",
}


def json_default(value: Any) -> str:
    if isinstance(value, (datetime, date)):
        return value.isoformat(sep=" ")
    return str(value)


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(data, ensure_ascii=False, indent=2, default=json_default),
        encoding="utf-8",
    )


def write_jsonl(path: Path, rows: Iterable[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as file:
        for row in rows:
            file.write(json.dumps(row, ensure_ascii=False, default=json_default) + "\n")


# ---------------------------------------------------------------------------
# 第一步：直接解析。这里尽量少做判断，让问题先暴露出来。
# ---------------------------------------------------------------------------


def parse_pdf(path: Path) -> dict[str, Any]:
    pages = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            pages.append(
                {
                    "page": page_number,
                    "text": page.extract_text() or "",
                    "tables": page.extract_tables() or [],
                }
            )
    return {"source": path.name, "type": "pdf", "pages": pages}


def iter_docx_blocks(document: Document) -> Iterable[Paragraph | Table]:
    """按正文真实顺序遍历段落和表格，避免把所有表格移动到文末。"""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def parse_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    blocks = []
    for block in iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            if block.text.strip():
                blocks.append(
                    {
                        "type": "paragraph",
                        "style": block.style.name if block.style else "",
                        "text": block.text.strip(),
                    }
                )
        else:
            blocks.append(
                {
                    "type": "table",
                    "rows": [
                        [cell.text.strip() for cell in row.cells]
                        for row in block.rows
                    ],
                }
            )

    section = document.sections[0]
    return {
        "source": path.name,
        "type": "docx",
        "header": [p.text for p in section.header.paragraphs if p.text.strip()],
        "footer": [p.text for p in section.footer.paragraphs if p.text.strip()],
        "blocks": blocks,
    }


def parse_xlsx(path: Path) -> dict[str, Any]:
    """同时读取公式版和值版，避免只保留公式或只保留显示值。"""
    formula_book = load_workbook(path, data_only=False, read_only=False)
    value_book = load_workbook(path, data_only=True, read_only=False)
    sheets = []

    for sheet_name in formula_book.sheetnames:
        formula_sheet = formula_book[sheet_name]
        value_sheet = value_book[sheet_name]
        rows = []
        for row_number in range(1, formula_sheet.max_row + 1):
            cells = []
            for column_number in range(1, formula_sheet.max_column + 1):
                formula_cell = formula_sheet.cell(row_number, column_number)
                value_cell = value_sheet.cell(row_number, column_number)
                cells.append(
                    {
                        "coordinate": formula_cell.coordinate,
                        "value": value_cell.value,
                        "formula": (
                            formula_cell.value
                            if formula_cell.data_type == "f"
                            else None
                        ),
                        "number_format": formula_cell.number_format,
                    }
                )
            rows.append(cells)
        sheets.append({"name": sheet_name, "rows": rows})

    return {"source": path.name, "type": "xlsx", "sheets": sheets}


# ---------------------------------------------------------------------------
# 第二步：清洗。不同格式使用不同规则，不做“一把梭”的字符串替换。
# ---------------------------------------------------------------------------


def edge_signature(line: str) -> str:
    """把页码等变化数字归一化，便于识别“第 1 页 / 第 2 页”。"""
    return re.sub(r"\d+", "{n}", line.strip())


def repeated_edge_lines(pages: list[dict[str, Any]], edge_size: int = 2) -> set[str]:
    """找出在多页顶部或底部重复出现的候选页眉页脚签名。"""
    candidates = []
    for page in pages:
        lines = [line.strip() for line in page["text"].splitlines() if line.strip()]
        candidates.extend(edge_signature(line) for line in lines[:edge_size])
        candidates.extend(edge_signature(line) for line in lines[-edge_size:])
    counts = Counter(candidates)
    return {line for line, count in counts.items() if count >= 2}


def merge_visual_line_breaks(lines: list[str]) -> list[str]:
    """合并 PDF 的视觉断行，但保留标题、句末和列表边界。"""
    merged: list[str] = []
    for line in lines:
        line = re.sub(r"\s+", " ", line).strip()
        if not line:
            continue
        is_heading = bool(re.match(r"^\d+(?:\.\d+)*\.\s*", line))
        previous_is_heading = bool(
            merged and re.match(r"^\d+(?:\.\d+)*\.\s*", merged[-1])
        )
        if (
            merged
            and not is_heading
            and not previous_is_heading
            and not re.search(r"[。！？；：]$", merged[-1])
            and len(merged[-1]) > 10
        ):
            separator = " " if re.search(r"[A-Za-z0-9]$", merged[-1]) else ""
            merged[-1] += separator + line
        else:
            merged.append(line)
    return merged


def markdown_table(rows: list[list[Any]]) -> list[str]:
    if not rows:
        return []
    normalized = [
        [str(cell or "").replace("\n", " ").strip() for cell in row]
        for row in rows
    ]
    return [
        "| " + " | ".join(normalized[0]) + " |",
        "| " + " | ".join(["---"] * len(normalized[0])) + " |",
        *["| " + " | ".join(row) + " |" for row in normalized[1:]],
    ]


def clean_pdf(parsed: dict[str, Any]) -> str:
    repeated = repeated_edge_lines(parsed["pages"])
    output = [f"# {Path(parsed['source']).stem}", ""]

    for page in parsed["pages"]:
        table_lines = {
            " ".join(str(cell or "").replace("\n", " ").strip() for cell in row)
            for table in page["tables"]
            for row in table
        }
        lines = [
            line.strip()
            for line in page["text"].splitlines()
            if edge_signature(line) not in repeated and line.strip() not in table_lines
        ]
        for line in merge_visual_line_breaks(lines):
            if re.match(r"^\d+(?:\.\d+)*\.\s*", line):
                output.extend([f"## {line}", ""])
            else:
                output.extend([line, ""])
        for index, table in enumerate(page["tables"], start=1):
            output.extend([f"**第 {page['page']} 页表格 {index}**", ""])
            output.extend(markdown_table(table))
            output.append("")
    return "\n".join(output).strip() + "\n"


def clean_docx(parsed: dict[str, Any]) -> str:
    output = [f"# {Path(parsed['source']).stem}", ""]
    for block in parsed["blocks"]:
        if block["type"] == "table":
            output.extend(markdown_table(block["rows"]))
            output.append("")
            continue

        text = block["text"]
        style = block.get("style", "")
        if style.startswith("Heading"):
            match = re.search(r"(\d+)$", style)
            level = min((int(match.group(1)) if match else 1) + 1, 6)
            output.extend(["#" * level + " " + text, ""])
        elif style.startswith("List Number"):
            output.extend(["1. " + text, ""])
        elif style.startswith("List Bullet"):
            output.extend(["- " + text, ""])
        else:
            output.extend([text, ""])
    return "\n".join(output).strip() + "\n"


def clean_xlsx(parsed: dict[str, Any]) -> list[dict[str, Any]]:
    records = []
    for sheet in parsed["sheets"]:
        header_index = None
        headers: list[str] = []
        required_headers = {"案例编号", "设备型号", "报警码"}
        for index, row in enumerate(sheet["rows"]):
            values = [cell["value"] for cell in row]
            value_set = {str(value).strip() for value in values if value not in (None, "")}
            if required_headers.issubset(value_set):
                header_index = index
                headers = [str(value or "").strip() for value in values]
                break
        if header_index is None:
            continue

        for row_index, row in enumerate(sheet["rows"][header_index + 1 :], start=header_index + 2):
            values = [cell["value"] for cell in row]
            if not any(value not in (None, "") for value in values):
                continue
            if values[0] in (None, ""):
                continue
            record = {
                "source": parsed["source"],
                "sheet": sheet["name"],
                "row": row_index,
            }
            for header, cell in zip(headers, row):
                if not header:
                    continue
                value = cell["value"]
                if header == "报警码" and value in (None, ""):
                    value = "未填写（不得推测）"
                elif value in (None, ""):
                    value = "未完成/未填写"
                record[header] = value
                if cell.get("formula"):
                    record[f"{header}_formula"] = cell["formula"]
            records.append(record)
    return records


# ---------------------------------------------------------------------------
# 第三步：切分。长文按标题，表格按记录；固定长度只用于对比。
# ---------------------------------------------------------------------------


def fixed_chunks(text: str, size: int = 180, overlap: int = 30) -> list[str]:
    plain = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0
    while start < len(plain):
        chunks.append(plain[start : start + size])
        start += size - overlap
    return chunks


def heading_chunks(markdown: str, source: str) -> list[dict[str, Any]]:
    title = Path(source).stem
    chunks = []
    current_section = "文档开头"
    buffer: list[str] = []

    def flush() -> None:
        content = "\n".join(buffer).strip()
        if content:
            chunks.append(
                {
                    "chunk_id": f"{Path(source).stem}-{len(chunks) + 1:03d}",
                    "source": source,
                    "title": title,
                    "section": current_section,
                    "content": content,
                }
            )

    for line in markdown.splitlines():
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if match:
            flush()
            buffer.clear()
            current_section = match.group(2).strip()
        else:
            buffer.append(line)
    flush()
    return chunks


def record_chunks(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    chunks = []
    for record in records:
        case_id = str(record.get("案例编号", "unknown"))
        content = "\n".join(
            f"{key}：{value}"
            for key, value in record.items()
            if key not in {"source", "sheet", "row"} and not key.endswith("_formula")
        )
        chunks.append(
            {
                "chunk_id": f"case-{case_id}",
                "source": record["source"],
                "sheet": record["sheet"],
                "row": record["row"],
                "content": content,
            }
        )
    return chunks


def run(input_dir: Path, output_dir: Path) -> None:
    paths = {key: input_dir / name for key, name in SAMPLE_FILES.items()}
    missing = [str(path) for path in paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少样例文件：\n" + "\n".join(missing))

    parsed = {
        "pdf": parse_pdf(paths["pdf"]),
        "docx": parse_docx(paths["docx"]),
        "xlsx": parse_xlsx(paths["xlsx"]),
    }
    for name, data in parsed.items():
        write_json(output_dir / "01_直接解析" / f"{name}.json", data)

    pdf_markdown = clean_pdf(parsed["pdf"])
    docx_markdown = clean_docx(parsed["docx"])
    xlsx_records = clean_xlsx(parsed["xlsx"])

    cleaned_dir = output_dir / "02_清洗后"
    cleaned_dir.mkdir(parents=True, exist_ok=True)
    (cleaned_dir / "pdf.md").write_text(pdf_markdown, encoding="utf-8")
    (cleaned_dir / "docx.md").write_text(docx_markdown, encoding="utf-8")
    write_json(cleaned_dir / "xlsx_records.json", xlsx_records)

    chunk_dir = output_dir / "03_切分后"
    write_jsonl(
        chunk_dir / "fixed_chunks.jsonl",
        [
            {"chunk_id": f"fixed-{index:03d}", "content": content}
            for index, content in enumerate(fixed_chunks(pdf_markdown), start=1)
        ],
    )
    write_jsonl(
        chunk_dir / "heading_chunks.jsonl",
        heading_chunks(pdf_markdown, paths["pdf"].name)
        + heading_chunks(docx_markdown, paths["docx"].name),
    )
    write_jsonl(chunk_dir / "record_chunks.jsonl", record_chunks(xlsx_records))

    print(f"处理完成：{output_dir.resolve()}")
    print("01_直接解析 → 02_清洗后 → 03_切分后")


def main() -> None:
    chapter_assets = Path(__file__).resolve().parent.parent
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=chapter_assets / "原始文件",
        help="PDF、DOCX、XLSX 样例所在目录",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path(__file__).resolve().parent / "运行结果",
        help="中间结果输出目录",
    )
    args = parser.parse_args()
    run(args.input_dir, args.output_dir)


if __name__ == "__main__":
    main()
